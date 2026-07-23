from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT_DIR = Path("artifacts")
BASE = "JOAT-Labs-pm-substrate-research-backed-thesis-and-arrowhedge-testbed"
DOCX_PATH = OUT_DIR / f"{BASE}.docx"
PDF_PATH = OUT_DIR / f"{BASE}.pdf"
MD_PATH = OUT_DIR / f"{BASE}.md"


TITLE = "State Coherence Under Partial Observation"
SUBTITLE = (
    "A research-backed rewrite of the JOAT Labs pm-substrate thesis, "
    "with ArrowHedgeLabs as the sandbox testbed"
)
BYLINE = "Prepared for Emmanuel Akinwale and JOAT Labs"
DATE_LINE = "May 26, 2026"


@dataclass
class Block:
    kind: str
    content: object


@dataclass
class Section:
    heading: str
    blocks: list[Block] = field(default_factory=list)


sections: list[Section] = [
    Section(
        "Executive Thesis",
        [
            Block(
                "p",
                "The modern workspace is not failing because teams lack software. "
                "It is failing because work now moves through many specialized "
                "tools, teams, and agents that each hold a partial model of the "
                "same changing reality. The hard problem is state coherence: "
                "how bounded actors can act safely when every actor sees only "
                "part of the system, every observation can be stale, and every "
                "action may change the state other actors depend on.",
            ),
            Block(
                "p",
                "JOAT Labs should frame pm-substrate as the project-manager layer "
                "for that problem. It is not another app and not a point-to-point "
                "integration hub. It is a tenant-scoped operational substrate that "
                "coordinates humans, tools, workflows, and AI agents through typed "
                "entities, relationships, events, capabilities, policies, workflow "
                "state, evidence, projections, commitments, and continuity records.",
            ),
            Block(
                "p",
                "The AI state problem is therefore not separate from the workspace "
                "gap. It is a more visible version of the same systems problem. "
                "An agent is a bounded perception-action system: it observes a "
                "partial world, encodes those observations into an internal model, "
                "chooses actions from that model, and receives delayed or incomplete "
                "feedback. Humans, departments, SaaS tools, and LLM agents all share "
                "this structure. The difference is that LLM agents make the failure "
                "observable because they act quickly across many systems without "
                "naturally owning institutional memory.",
            ),
            Block(
                "callout",
                "Rewritten claim: pm-substrate does not solve reality. It solves "
                "governed operational state. It records what the workspace accepts "
                "as actionable, where that claim came from, how fresh it is, which "
                "actor may change it, which workflow it belongs to, and how agents "
                "can resume from it after context loss.",
            ),
        ],
    ),
    Section(
        "What Changed From The Earlier Thesis",
        [
            Block(
                "p",
                "The earlier draft was directionally strong because it identified "
                "state as the hidden bottleneck behind both AI agents and cross-tool "
                "project management. The rewrite sharpens that claim in four ways.",
            ),
            Block(
                "bullets",
                [
                    "First, it separates reality state, observed state, belief state, operational state, authoritative state, historical state, and projection state. This avoids claiming that the substrate knows reality directly.",
                    "Second, it makes AI a semantic I/O and mapping assistant, not the source of authority. The LLM proposes mappings and transformations; deterministic validators, profiles, permissions, and write gates decide admissibility.",
                    "Third, it makes plug-in adoption measurable. A platform sits on top of pm-substrate when it can map into the substrate with new mapping/profile/capability files, not by rewriting the substrate or hard-coding itself into existing providers.",
                    "Fourth, it moves ArrowHedgeLabs from an example to a sandbox. The financial research agents already expose the exact problems the substrate should solve: many specialist agents, shared state, risk constraints, workflow edges, evidence, replay, and amnesiac continuity.",
                ],
            ),
        ],
    ),
    Section(
        "The First-Principles Problem",
        [
            Block(
                "p",
                "At the lowest useful level, an agent is not a chatbot. It is a "
                "stateful system coupled to an environment. The environment has "
                "state S. The agent receives observations O, maintains an internal "
                "model M, updates that model as observations arrive, chooses an "
                "action A through a policy or objective, and then receives feedback "
                "from the changed environment. This structure appears in control "
                "systems, animals, humans, departments, software services, robots, "
                "and LLM agents.",
            ),
            Block(
                "p",
                "Once the agent is decomposed, the state problem appears at every "
                "layer: observation can be partial, representation can lose "
                "structure, memory can become stale, goals can conflict, plans can "
                "be invalidated by new events, tools can mutate reality, feedback "
                "can be delayed, and identity can disappear across sessions. The "
                "LLM context window is only one surface of this deeper problem.",
            ),
            Block(
                "p",
                "Physics is useful here as a discipline of humility. Observation "
                "does not give unlimited, undistorted access to a system. In "
                "operational systems, the analogy is not quantum mechanics itself; "
                "the analogy is the discipline of admitting that measurement is "
                "partial, instruments have side effects, and a system's state must "
                "be estimated from bounded observations. pm-substrate should not "
                "pretend to hold perfect truth. It should make observation, evidence, "
                "uncertainty, freshness, authority, and invalidation explicit.",
            ),
        ],
    ),
    Section(
        "Research Findings That Support The Approach",
        [
            Block(
                "p",
                "The thesis is strongest when grounded in work from several fields. "
                "The Arrowsmith or literature-based discovery move is to look for "
                "nearby domains that solved pieces of the same abstract problem, "
                "then compose the useful mechanisms without pretending any one "
                "domain solved the whole workspace.",
            ),
            Block(
                "table",
                [
                    ["Finding", "Research or standard", "Design consequence for pm-substrate"],
                    [
                        "Autonomous agents require profile, memory, planning, and action components.",
                        "Wang et al. survey LLM-based autonomous agents around profile, memory, planning, and action [R1].",
                        "Agent memory must be a substrate layer above events and workflows, not just prompt history.",
                    ],
                    [
                        "Complex LLM workflows benefit from explicit states and transitions.",
                        "StateFlow proposes state-driven workflows for LLM task solving [R2].",
                        "pm-substrate should expose workflows as durable state machines with valid transitions.",
                    ],
                    [
                        "Multi-agent systems need infrastructure for interaction, memory, and coordination.",
                        "Li et al. survey LLM multi-agent systems through workflow, infrastructure, and challenges [R3].",
                        "Agents should coordinate through shared substrate events and capabilities, not hidden chat.",
                    ],
                    [
                        "Cross-functional performance depends on inter-team coordination and information elaboration.",
                        "Xie et al. link coordination, information elaboration, and team performance [R4].",
                        "The substrate's business value is reduced coordination cost, not isolated intelligence.",
                    ],
                    [
                        "Knowledge workers switch applications because work is distributed across task surfaces.",
                        "Jahanlou et al. characterize task-centric application switching in knowledge work [R5].",
                        "The workspace gap is real at the task level; the system must preserve context across tools.",
                    ],
                    [
                        "Provenance can be modeled as entities, activities, and agents.",
                        "W3C PROV defines a standard provenance model [S1].",
                        "Every state claim should link to source, actor, activity, and evidence.",
                    ],
                    [
                        "Declarative mappings let existing schemas project into another model.",
                        "W3C R2RML maps relational databases to RDF using mapping documents [S2].",
                        "pm-substrate should use declarative entity mappings rather than app rewrites.",
                    ],
                    [
                        "JSON data contracts can be validated with formal schemas.",
                        "JSON Schema defines a declarative vocabulary for annotating and validating JSON [S3].",
                        "AI outputs and platform mappings need schema-constrained proposals and validation gates.",
                    ],
                    [
                        "Interoperable data needs language-neutral formats.",
                        "Apache Arrow specifies a language-independent columnar memory format [S4].",
                        "The substrate should treat formats as adapters and preserve logical meaning above file form.",
                    ],
                    [
                        "Controllers reconcile desired state and current state in loops.",
                        "Kubernetes controllers watch state and move current state toward desired state [S5].",
                        "Cron jobs and heartbeats should be reconciliation and freshness checks, not the source of truth.",
                    ],
                    [
                        "Risk-managed AI requires measurement, governance, and controls.",
                        "NIST AI RMF frames AI risk management through govern, map, measure, and manage functions [S6].",
                        "The substrate needs measurable risks, authority boundaries, and audit-ready controls.",
                    ],
                    [
                        "LLMs can produce schema-constrained outputs, but the schema is the control.",
                        "OpenAI Structured Outputs constrain model output to developer-supplied JSON Schema [S7].",
                        "AI can propose plug-in mappings, but accepted substrate writes must pass deterministic contracts.",
                    ],
                ],
            ),
            Block(
                "p",
                "The common pattern is clear. Other disciplines do not say, 'make "
                "one giant intelligent thing.' They externalize state into ledgers, "
                "models, schemas, events, workflows, controllers, provenance records, "
                "and audit trails. pm-substrate is the same pattern applied to "
                "agentic workspaces.",
            ),
        ],
    ),
    Section(
        "The Key Distinction: AI As Semantic I/O, Not Authority",
        [
            Block(
                "p",
                "The user's intuition is right: modern AI is unusually good at "
                "turning many kinds of input into structured output. Text, images, "
                "audio, tables, CSV files, JSON, database schemas, API descriptions, "
                "and messy documents can all be interpreted by capable models. "
                "That makes AI the best available semantic adapter layer.",
            ),
            Block(
                "p",
                "But the substrate should not conclude that AI is the schema. The "
                "safe architecture is compiler-like. AI reads the source platform's "
                "shape and proposes a mapping. Deterministic validators check the "
                "mapping against a profile. Dry-run ingestion tests the mapping "
                "against sample rows and events. Human or policy approval promotes "
                "the mapping version. Runtime adapters apply the mapping. Substrate "
                "write gates validate every mutation. Event provenance records what "
                "changed and why.",
            ),
            Block(
                "callout",
                "Formula: AI is the semantic mapper. pm-substrate is the type system, compiler, runtime, event ledger, and audit trail.",
            ),
            Block(
                "p",
                "This avoids the two-state problem. There should not be one "
                "independent AI state and one independent project state competing "
                "for authority. Agent state should be interpretive and evidence-backed. "
                "Project state should be governed and authoritative only after "
                "validation. Derived views should carry freshness metadata. When "
                "they disagree, the substrate does not ask which story sounds better; "
                "it follows source authority, event history, profile rules, and "
                "conflict procedures.",
            ),
        ],
    ),
    Section(
        "How A Platform Plugs In Without Being Rewritten",
        [
            Block(
                "p",
                "The implementation baseline already exists in the pm-substrate "
                "repository. ADR-0020 defines the declarative entity-mapping format. "
                "ADR-0021 adds profile-aware semantic validation. ADR-0022 adds an "
                "ingestion adapter that turns a source row and tenant context into "
                "graph-ready inputs. That is the start of the easy plug-in story.",
            ),
            Block(
                "numbered",
                [
                    "Discover the source. Inspect the platform's database schema, API schema, CSV headers, event payloads, or sample exports.",
                    "Ask AI to propose the mapping. The model proposes entity mappings, field aliases, edges, source references, event types, and capability descriptors.",
                    "Validate structurally. The mapping must match the substrate mapping format: version, profile, entity declarations, identity fields, field maps, schema versions, and edge declarations.",
                    "Validate semantically. The mapping must resolve against the installed profile: concrete type exists, tier-1 binding matches, required identity fields are covered, edge types and cardinalities match.",
                    "Dry-run sample data. Apply the mapping to representative rows and events. The adapter returns graph-ready node and edge inputs without mutating production state.",
                    "Run write-gate tests. The graph, event log, registry, permissions, workflow runtime, and provenance rules accept or reject the proposed writes deterministically.",
                    "Version and approve. Store the mapping version, sample fixtures, expected outputs, provenance policy, and rollback path.",
                    "Operate through adapters. The source platform remains itself. It only needs an export, API, webhook, CDC stream, or small adapter. The substrate is not rewritten to fit it.",
                ],
            ),
            Block(
                "p",
                "The acceptance criterion is concrete: onboarding a new platform "
                "should require new mapping files, fixtures, tests, and maybe a thin "
                "adapter, but zero changes to substrate packages and zero changes "
                "to existing providers. If a new platform requires changing the "
                "graph, event log, registry, workflow runtime, or another provider, "
                "the plug-in claim failed.",
            ),
        ],
    ),
    Section(
        "Dynamic And Structurally Correct",
        [
            Block(
                "p",
                "Dynamic does not mean schema-less. It means the system can accept "
                "new source shapes while preserving explicit contracts. A platform "
                "can bring CSV, JSON, SQL rows, API payloads, documents, or event "
                "streams. AI can infer what each field likely means, but the substrate "
                "only accepts a mapping once it compiles against a known profile and "
                "passes sample-based tests.",
            ),
            Block(
                "table",
                [
                    ["Layer", "Dynamic behavior", "Correctness control"],
                    ["File and API intake", "Accept CSV, JSON, SQL, webhooks, APIs, documents, and multimodal evidence.", "Use parsers, MIME/type detection, checksums, sample fixtures, and source metadata."],
                    ["Semantic interpretation", "AI proposes entities, relationships, field aliases, and events.", "Use structured outputs, mapping schemas, confidence thresholds, review queues, and rejected-output logs."],
                    ["Profile binding", "Different vertical profiles can define concrete types over shared Tier-1 primitives.", "Run profile-aware semantic validation before any write path."],
                    ["Runtime ingestion", "Adapters transform source records into graph/event inputs.", "Write gates enforce tenant, identity, schema version, permissions, idempotency, and provenance."],
                    ["State propagation", "Events update projections, workflows, and agent continuity.", "Replay tests, causation chains, freshness metadata, and contradiction detection."],
                    ["Evolution", "Mappings can change as source platforms change.", "Version mappings, migration plans, backfills, compatibility tests, and rollback paths."],
                ],
            ),
            Block(
                "p",
                "This is how the substrate stays both flexible and structurally "
                "correct. AI handles semantic interpolation. The substrate handles "
                "contracts. That split is the architecture.",
            ),
        ],
    ),
    Section(
        "Heartbeats, Cron Jobs, And The Observation Problem",
        [
            Block(
                "p",
                "The periodic heartbeat idea is useful, but it should be framed "
                "carefully. A cron job should not be treated as truth. It should be "
                "treated as a reconciliation event: a scheduled opportunity to "
                "observe source systems, refresh projections, detect drift, expire "
                "stale claims, and force agents to rebase their local plans against "
                "the current substrate.",
            ),
            Block(
                "p",
                "Borrow the Kubernetes controller pattern: watch desired state and "
                "current state, then reconcile differences. Borrow observability "
                "discipline from OpenTelemetry: traces, metrics, and logs are separate "
                "signals that help explain system behavior. Borrow distributed-systems "
                "discipline from Lamport: ordered events and causal relationships "
                "matter when components do not share a single clock.",
            ),
            Block(
                "p",
                "In pm-substrate, a heartbeat should emit a state census. The census "
                "does not claim perfect reality. It says: these sources were checked, "
                "these facts were fresh, these projections were invalidated, these "
                "workflows were blocked, these agents have open plans depending on "
                "changed state, and these conflicts need human or policy resolution.",
            ),
        ],
    ),
    Section(
        "Why ArrowHedgeLabs Is The Right Sandbox",
        [
            Block(
                "p",
                "ArrowHedgeLabs is a strong sandbox because it already looks like a "
                "multi-agent organization. The local project contains multiple "
                "specialist analyst agents, a risk manager, a portfolio manager, a "
                "LangGraph workflow, structured state, and a backtesting orientation. "
                "It also contains Dexter, a research-agent environment with scratchpad, "
                "memory, compaction, and tool execution. That makes it close to the "
                "problem pm-substrate claims to solve.",
            ),
            Block(
                "table",
                [
                    ["Sandbox feature", "Why it matters"],
                    ["Many specialist analysts", "Represents cross-functional departments with different beliefs and evidence."],
                    ["Risk manager", "Represents deterministic constraints that should gate agent proposals."],
                    ["Portfolio manager", "Already separates LLM judgment from deterministic allowed-action computation."],
                    ["LangGraph state", "Shows current local state shape: messages, data, metadata, and merged analyst signals."],
                    ["Backtesting", "Provides repeatable historical scenarios for replay and regression testing."],
                    ["Dexter scratchpad and memory", "Shows the difference between local agent notes and substrate-governed continuity."],
                ],
            ),
            Block(
                "p",
                "This sandbox must remain research and education only, not financial "
                "advice or trading infrastructure. Its value is that financial "
                "research has rich state: tickers, dates, evidence, analyst signals, "
                "risk limits, portfolio constraints, decisions, outcomes, and replay. "
                "That richness makes it a demanding testbed for state coherence.",
            ),
        ],
    ),
    Section(
        "ArrowHedgeLabs Test Plan",
        [
            Block(
                "p",
                "The testbed should validate two claims at the same time: first, "
                "that existing platforms can plug into the substrate through mappings "
                "and adapters; second, that agents behave better when they resume "
                "from substrate state rather than local chat history.",
            ),
            Block(
                "table",
                [
                    ["Test", "Procedure", "Pass condition"],
                    ["T1: Source mapping", "Map ArrowHedgeLabs tickers, portfolio, analyst signals, risk outputs, and decisions into substrate entities/events.", "No substrate package edits; mapping validates structurally and semantically."],
                    ["T2: Event provenance", "Each analyst signal becomes an event with source, timestamp, evidence payload, ticker, model/provider metadata, and causation chain.", "Every final decision can trace back to its contributing signals and risk state."],
                    ["T3: Deterministic risk gate", "Keep compute_allowed_actions as a deterministic boundary and record its output as validated operational state.", "LLM can choose only actions and quantities permitted by the risk gate."],
                    ["T4: Amnesiac resume", "Run session 1, store checkpoints, delete chat context, then run session 2 from tenant, agent, and scope only.", "Agent resumes open work, avoids contradicted claims, and cites substrate evidence."],
                    ["T5: Staleness failure", "Replay a case where price or position state changes after analyst signals are created.", "Workflow blocks or requests refresh before producing a stale decision."],
                    ["T6: Plug-in file format", "Ingest the same scenario from JSON, CSV, and a SQL-like row export.", "All formats produce equivalent canonical substrate entities/events through mappings."],
                    ["T7: Replay audit", "Replay a historical backtest from event history and compare projected portfolio decisions.", "Replay reproduces decisions or explains deterministic differences from changed mappings/profiles."],
                    ["T8: Conflict handling", "Create conflicting analyst claims and stale risk data.", "System records conflict, prevents silent promotion, and routes to rule or human review."],
                ],
            ),
        ],
    ),
    Section(
        "Metrics To Monitor",
        [
            Block(
                "p",
                "The substrate should be judged by behavior, not elegance. These "
                "metrics tell whether it is solving state coherence and plug-in "
                "adoption rather than simply storing more data.",
            ),
            Block(
                "table",
                [
                    ["Metric", "Definition"],
                    ["Time-to-plugin", "Elapsed time from first source schema/sample to validated mapping, dry-run, and first accepted substrate write."],
                    ["Substrate edit count", "Number of changes required in substrate packages to onboard a new platform. Target: zero."],
                    ["Mapping coverage", "Percent of source entities, events, and required fields mapped to Tier-1/profile concepts."],
                    ["Validator rejection rate", "Percent of AI-proposed mappings rejected by structural, semantic, or sample-based validation."],
                    ["Evidence coverage", "Percent of authoritative state transitions linked to source evidence and actor/activity provenance."],
                    ["State disagreement rate", "Frequency of conflicts between source systems, projections, and agent beliefs."],
                    ["Stale action rate", "Percent of attempted actions based on expired or invalidated state."],
                    ["Agent resume success", "Percent of amnesiac runs that resume correct scope, open work, constraints, and evidence without chat history."],
                    ["Replay fidelity", "Percent of historical workflows whose outputs can be reproduced or explained from event history."],
                    ["Unauthorized action block rate", "Blocked attempts that would have mutated state without permission or valid workflow position."],
                    ["Cross-tool outcome success", "Percent of workflows completed across multiple providers without provider-to-provider hard coupling."],
                    ["Mean time to reconcile", "Elapsed time from conflict or drift detection to accepted resolution."],
                ],
            ),
        ],
    ),
    Section(
        "Risks And Mitigation Strategy",
        [
            Block(
                "table",
                [
                    ["Risk", "Why it matters", "Mitigation"],
                    ["AI hallucinated mapping", "The model may infer a field meaning incorrectly.", "Use schema-constrained outputs, sample fixtures, deterministic validators, human approval for promotion, and regression tests."],
                    ["Two-state problem", "Agent memory can diverge from project state.", "Make agent memory subordinate: claims must cite evidence and rebase against substrate state before action."],
                    ["Source authority ambiguity", "Two tools may disagree about the same fact.", "Maintain source-of-truth registry by fact type and conflict resolution policy."],
                    ["Schema drift", "Source platforms change fields or semantics.", "Version mappings, monitor ingestion errors, run nightly schema diff, and require migration plans."],
                    ["Over-instrumentation", "Observation itself can change behavior or create noise.", "Separate telemetry from authority and use heartbeats as reconciliation, not truth."],
                    ["Workflow brittleness", "Hard-coded flows break when capabilities change.", "Use capability descriptors, workflow soundness checks, optional providers, and dead-letter paths."],
                    ["Privacy and tenant leakage", "Cross-tenant operational state is sensitive.", "Tenant-scoped writes, permission gates, provenance, audit logs, and least-privilege capability registration."],
                    ["Financial misuse in sandbox", "A trading-like demo can be mistaken for investment advice.", "Keep ArrowHedgeLabs as historical/research simulation with no real-trading path."],
                ],
            ),
        ],
    ),
    Section(
        "Implementation Roadmap",
        [
            Block(
                "numbered",
                [
                    "Define the ArrowHedgeLabs research profile: Ticker, ResearchRun, AnalystSignal, RiskState, PortfolioState, Decision, EvidenceDocument, and BacktestRun over existing Tier-1 primitives.",
                    "Write the first ArrowHedgeLabs entity mapping and fixtures. Validate it using the existing structural and semantic validators.",
                    "Add event contracts for analyst.signal.created, risk.state.validated, portfolio.decision.proposed, portfolio.decision.accepted, and workflow.blocked.stale_state.",
                    "Build a read-only ingestion adapter that projects one historical ArrowHedgeLabs run into the substrate without mutating the source project.",
                    "Record deterministic risk-gate output as operational state and require decisions to cite it.",
                    "Store continuity checkpoints for analyst conclusions, risk decisions, unresolved questions, and handoffs.",
                    "Run the amnesiac-agent evaluation using ArrowHedgeLabs: delete chat history and force the agent to resume from substrate continuity plus event history.",
                    "Add the plug-in acceptance test: onboarding ArrowHedgeLabs must not modify substrate packages or existing providers.",
                ],
            ),
            Block(
                "p",
                "This roadmap validates the original PM-layer baseline. The "
                "substrate first proves it can coordinate independent providers "
                "through mappings, capabilities, events, and workflows. Then the "
                "agentic layer proves the same coordination substrate improves "
                "agent continuity, state freshness, and evidence-backed action. "
                "The agentic state thesis is therefore not a pivot. It is a harder "
                "validation surface for the original project-manager-layer thesis.",
            ),
        ],
    ),
    Section(
        "Conclusion",
        [
            Block(
                "p",
                "The pm-substrate thesis becomes strongest when stated simply: "
                "work fails when bounded actors act from divergent local state; "
                "the project-manager layer is the governed operational substrate "
                "that keeps those actors aligned. AI makes the opportunity larger "
                "because it can translate across messy inputs, formats, and schemas. "
                "AI also makes the risk larger because it can act quickly from "
                "partial or stale state.",
            ),
            Block(
                "p",
                "The solution is not to make the model the source of truth. The "
                "solution is to make the substrate the governed state plane and "
                "use AI as the semantic adapter into it. That is how a platform "
                "can sit on top of pm-substrate without being rewritten: map it, "
                "validate it, dry-run it, version it, and let the substrate enforce "
                "the runtime contract.",
            ),
            Block(
                "p",
                "ArrowHedgeLabs is the right sandbox because it is already a "
                "miniature cross-functional organization of agents, constraints, "
                "evidence, workflows, and decisions. If pm-substrate can make that "
                "system replayable, evidence-backed, resumable, and plug-in clean "
                "without rewriting the substrate, it will have validated the core "
                "JOAT Labs claim: design the interactions first.",
            ),
        ],
    ),
]


references = [
    ("R1", "Wang, L. et al. (2024). A survey on large language model based autonomous agents. Frontiers of Computer Science. https://link.springer.com/article/10.1007/s11704-024-40231-1"),
    ("R2", "Wu, Y. et al. (2024). StateFlow: Enhancing LLM Task-Solving through State-Driven Workflows. COLM 2024. https://openreview.net/forum?id=3nTbuygoop"),
    ("R3", "Li, X. et al. (2024). A survey on LLM-based multi-agent systems: workflow, infrastructure, and challenges. Vicinagearth. https://link.springer.com/article/10.1007/s44336-024-00009-2"),
    ("R4", "Xie, X., Ling, C.-D., Liu, W., & Wei, J. (2022). Inter-team coordination, information elaboration, and performance in teams. Journal of Business Research. https://doi.org/10.1016/j.jbusres.2022.05.002"),
    ("R5", "Jahanlou, A. et al. (2023). Task-Centric Application Switching. Graphics Interface. https://www.research.autodesk.com/publications/task-centric-application-switching/"),
    ("R6", "Hatalis, K. et al. (2024). Memory Matters: The Need to Improve Long-Term Memory in LLM-Agents. AAAI Symposium. https://ojs.aaai.org/index.php/AAAI-SS/article/view/27688"),
    ("R7", "Zhang, Y. et al. (2023). Schema Matching using Pre-Trained Language Models. IEEE ICDE. https://www.microsoft.com/en-us/research/publication/schema-matching-using-pre-trained-language-models/"),
    ("R8", "Sebastian, A. et al. (2021). A systematic review of literature-based discovery. https://pmc.ncbi.nlm.nih.gov/articles/PMC7924697/"),
    ("S1", "W3C. PROV Overview. https://www.w3.org/TR/prov-overview/"),
    ("S2", "W3C. R2RML: RDB to RDF Mapping Language. https://www.w3.org/TR/r2rml/"),
    ("S3", "JSON Schema. What is JSON Schema? https://json-schema.org/overview/what-is-jsonschema"),
    ("S4", "Apache Arrow. Columnar format and language-independent data. https://arrow.apache.org/docs/format/Intro.html"),
    ("S5", "Kubernetes. Controllers. https://kubernetes.io/docs/concepts/architecture/controller/"),
    ("S6", "NIST. AI Risk Management Framework. https://www.nist.gov/itl/ai-risk-management-framework"),
    ("S7", "OpenAI. Structured Outputs. https://platform.openai.com/docs/guides/structured-outputs"),
    ("S8", "OpenTelemetry. What is OpenTelemetry? https://opentelemetry.io/docs/what-is-opentelemetry/"),
    ("S9", "Lamport, L. (1978). Time, Clocks, and the Ordering of Events in a Distributed System. https://www.microsoft.com/en-us/research/publication/time-clocks-ordering-events-distributed-system/"),
    ("S10", "Git. What is Git? https://git-scm.com/book/en/v2/Getting-Started-What-is-Git%3F"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(SUBTITLE)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("4B5563")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run(BYLINE).italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(DATE_LINE)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Purpose: replace the thesis draft with a research-backed thesis, "
        "implementation test plan, plug-in strategy, risk register, and sandbox plan."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("4B5563")
    doc.add_page_break()


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    col_count = len(rows[0])
    if col_count == 2:
        widths = [2600, 6760]
    elif col_count == 3:
        widths = [2400, 3300, 3660]
    else:
        widths = [9360 // col_count for _ in range(col_count)]
    set_table_width(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            run.font.size = Pt(9)
            if row_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("1F3A5F")
    doc.add_paragraph()


def add_docx_block(doc: Document, block: Block) -> None:
    if block.kind == "p":
        doc.add_paragraph(str(block.content))
    elif block.kind == "callout":
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        set_table_width(table, [9360])
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F4F6F9")
        set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(str(block.content))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("1F3A5F")
        doc.add_paragraph()
    elif block.kind == "bullets":
        for item in block.content:  # type: ignore[union-attr]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(str(item))
    elif block.kind == "numbered":
        for item in block.content:  # type: ignore[union-attr]
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(str(item))
    elif block.kind == "table":
        add_docx_table(doc, block.content)  # type: ignore[arg-type]
    else:
        raise ValueError(f"Unknown block kind: {block.kind}")


def build_docx() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)
    for section in sections:
        doc.add_heading(section.heading, level=1)
        for block in section.blocks:
            add_docx_block(doc, block)

    doc.add_page_break()
    doc.add_heading("References", level=1)
    for key, ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"[{key}] ")
        r.bold = True
        p.add_run(ref)

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(DOCX_PATH)


def pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "title",
            parent=base["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=base["BodyText"],
            alignment=TA_CENTER,
            fontName="Helvetica",
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=14,
        ),
        "meta": ParagraphStyle(
            "meta",
            parent=base["BodyText"],
            alignment=TA_CENTER,
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=6,
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "cell": ParagraphStyle(
            "cell",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.3,
            leading=9,
            spaceAfter=0,
        ),
        "cell_header": ParagraphStyle(
            "cell_header",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=7.5,
            leading=9,
            textColor=colors.HexColor("#1F3A5F"),
            spaceAfter=0,
        ),
        "callout": ParagraphStyle(
            "callout",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.4,
            leading=13,
            textColor=colors.HexColor("#1F3A5F"),
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#D0D7DE"),
            borderWidth=0.5,
            borderPadding=8,
            spaceBefore=4,
            spaceAfter=10,
        ),
    }
    return styles


def add_pdf_table(story, rows: list[list[str]], styles) -> None:
    col_count = len(rows[0])
    if col_count == 2:
        col_widths = [1.8 * inch, 4.7 * inch]
    elif col_count == 3:
        col_widths = [1.5 * inch, 2.25 * inch, 2.75 * inch]
    else:
        col_widths = [6.5 * inch / col_count] * col_count
    table_data = []
    for row_idx, row in enumerate(rows):
        row_cells = []
        for cell in row:
            style = styles["cell_header"] if row_idx == 0 else styles["cell"]
            row_cells.append(Paragraph(str(cell), style))
        table_data.append(row_cells)
    table = Table(table_data, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def add_pdf_block(story, block: Block, styles) -> None:
    if block.kind == "p":
        story.append(Paragraph(str(block.content), styles["body"]))
    elif block.kind == "callout":
        story.append(Paragraph(str(block.content), styles["callout"]))
    elif block.kind == "bullets":
        items = [
            ListItem(Paragraph(str(item), styles["body"]), leftIndent=14)
            for item in block.content  # type: ignore[union-attr]
        ]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=16))
        story.append(Spacer(1, 4))
    elif block.kind == "numbered":
        items = [
            ListItem(Paragraph(str(item), styles["body"]), leftIndent=14)
            for item in block.content  # type: ignore[union-attr]
        ]
        story.append(ListFlowable(items, bulletType="1", leftIndent=16))
        story.append(Spacer(1, 4))
    elif block.kind == "table":
        add_pdf_table(story, block.content, styles)  # type: ignore[arg-type]
    else:
        raise ValueError(f"Unknown block kind: {block.kind}")


def build_pdf() -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=TITLE,
        author="JOAT Labs",
    )
    story = [
        Paragraph(TITLE, styles["title"]),
        Paragraph(SUBTITLE, styles["subtitle"]),
        Paragraph(BYLINE, styles["meta"]),
        Paragraph(DATE_LINE, styles["meta"]),
        Spacer(1, 16),
        Paragraph(
            "Purpose: replace the thesis draft with a research-backed thesis, "
            "implementation test plan, plug-in strategy, risk register, and sandbox plan.",
            styles["meta"],
        ),
        PageBreak(),
    ]
    for section in sections:
        story.append(Paragraph(section.heading, styles["h1"]))
        for block in section.blocks:
            add_pdf_block(story, block, styles)
    story.append(PageBreak())
    story.append(Paragraph("References", styles["h1"]))
    for key, ref in references:
        story.append(Paragraph(f"<b>[{key}]</b> {ref}", styles["body"]))
    doc.build(story)


def markdown_escape(text: str) -> str:
    return text.replace("\n", " ")


def build_markdown() -> None:
    lines: list[str] = [
        f"# {TITLE}",
        "",
        SUBTITLE,
        "",
        BYLINE,
        "",
        DATE_LINE,
        "",
    ]
    for section in sections:
        lines.extend([f"## {section.heading}", ""])
        for block in section.blocks:
            if block.kind == "p":
                lines.extend([markdown_escape(str(block.content)), ""])
            elif block.kind == "callout":
                lines.extend([f"> **Key point:** {markdown_escape(str(block.content))}", ""])
            elif block.kind == "bullets":
                for item in block.content:  # type: ignore[union-attr]
                    lines.append(f"- {markdown_escape(str(item))}")
                lines.append("")
            elif block.kind == "numbered":
                for idx, item in enumerate(block.content, 1):  # type: ignore[union-attr]
                    lines.append(f"{idx}. {markdown_escape(str(item))}")
                lines.append("")
            elif block.kind == "table":
                rows = block.content  # type: ignore[assignment]
                header = rows[0]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(markdown_escape(str(cell)) for cell in row) + " |")
                lines.append("")
    lines.extend(["## References", ""])
    for key, ref in references:
        lines.extend([f"- [{key}] {ref}", ""])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    build_markdown()
    build_docx()
    build_pdf()
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")


if __name__ == "__main__":
    main()
